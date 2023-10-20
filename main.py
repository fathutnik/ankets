import docx
from docx.shared import Pt


filepath = r'/Users/nikita/PycharmProjects/Anketa/Анкета2.docx'
doc = docx.Document(filepath)

table = doc.tables[0]

data = []

for table in doc.tables:
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text not in row_data:
                    row_data.append(cell_text)
            data.append(row_data)

name, address, ogrn, inn, account = '', '', '', '', ''

for i in range(len(data)):
    for j in range(len(data[i])):
        if 'ОРГАНИЗАЦИЯ ' in data[i][j]:
            name = data[i+1][j].replace('Полное наименование организации: ', '')
        if 'ОРГАНИЗАЦИЯ' in data[i][j]:
            details = data[i+1][j+1].split('\n')
            inn = details[0].replace('ИНН: ', '')
            ogrn = details[-1].replace('ОГРН: ', '')
        if 'Адрес почтовый' in data[i][j]:
            address = data[i][j].replace('Адрес почтовый (с индексом): ', '').split('\n')[0]
        if 'р/с' in data[i][j]:
            if 'рублях' in data[i][j]:
                account = data[i][j].split('\n')[1].split(' ')[1]
            if 'юани' in data[i][j]:
                account = data[i][j].split('\n')[1].split(' ')[1]

email = data[6][0].split('\n')[0].split(' ')[2].replace('>;', '')
number = data[6][1].split('\n')[0].split(' ')[2]

filepath = r'/Users/nikita/PycharmProjects/Anketa/Заявление.docx'
doc = docx.Document(filepath)

doc.paragraphs[2].text = name

table = doc.tables[0]

table.cell(1, 1).text = name
table.cell(2, 1).text = address
table.cell(4, -1).text = ogrn
table.cell(5, -1).text = inn
table.cell(6, -1).text = number
table.cell(7, -1).text = account

for row in table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)
                paragraph.alignment = 1

table = doc.tables[1]
table.cell(1, -1).text = email

for row in table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)
                paragraph.alignment = 1

doc.save(r'/Users/nikita/PycharmProjects/Anketa/Заявление ' + name + '.docx')


